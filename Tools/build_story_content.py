#!/usr/bin/env python3
"""Convert the V2 Word screenplay into deterministic, validated game JSON.

The Word document remains the editorial source of truth.  This builder keeps every
player-facing screenplay paragraph verbatim, turns marked choices into branches,
and embeds stable source paragraph IDs so content can be rebuilt without hand edits.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "《夏末放映室》剧本V2.docx"
OUTPUT = ROOT / "GameTemplate" / "GameTemplate" / "Resources" / "StoryContent.json"

START_INDEX = 56
STOP_INDEX = 1410

BACKGROUND_RULES = [
    ("回到潮见坂", "environment_coastal_train_summer_afternoon"),
    ("旧宅与信箱", "environment_sahara_house_twilight"),
    ("胶片中的自己", "environment_lu_room_late_night"),
    ("最后的电影部", "environment_final_film_club_room"),
    ("《九月的第一天》", "environment_final_film_club_room"),
    ("第一次放映", "environment_underground_projection_room"),
    ("剧本预言", "environment_final_film_club_room"),
    ("黄昏里的少女", "environment_old_auditorium_twilight"),
    ("悠真的录音带", "environment_old_broadcast_room"),
    ("被剪开的照片", "environment_library_archive_corner"),
    ("灯里的谎言", "environment_school_rooftop_sunset"),
    ("不要让他们想起来", "environment_lu_room_late_night"),
    ("道具箱里的蓝色", "environment_final_film_club_room"),
    ("海鸥电影院", "environment_seagull_cinema_exterior"),
    ("五年前的事故", "environment_flooded_cinema_accident"),
    ("忘潮铃", "environment_seaside_shrine_bell"),
    ("最后一次放映", "environment_private_final_screening"),
    ("只留下夏天", "environment_private_final_screening"),
    ("没有第五个人", "environment_seagull_cinema_projection_room"),
    # This ending opens in the private screening room and only changes to the
    # coastal train once Riku leaves for Shirahama the following morning.
    ("海的另一边", "environment_private_final_screening"),
    ("放映结束以后", "environment_private_final_screening"),
    ("写给九月", "environment_rooftop_september_morning"),
]

BGM_BY_CHAPTER = {
    "序章": "M02_back_to_chomi_zaka",
    "第一章": "M03_film_club_daily",
    "第二章": "M05_chinatsu_theme",
    "第三章": "M04_investigation_clues",
    "第四章": "M08_typhoon_night",
    "第五章": "M10_final_screening",
    "只留下夏天": "E01_summer_only",
    "没有第五个人": "E02_no_fifth_person",
    "海的另一边": "E03_other_side_of_sea",
    "放映结束以后": "E04_after_screening",
    "写给九月": "E05_letter_to_september",
}

PORTRAITS = {
    "灯里": "akari_talking",
    "雨宫灯里": "akari_talking",
    "陆": "riku_neutral",
    "佐原陆": "riku_neutral",
    "少年陆": "riku_neutral",
    "诗织": "portrait_shiori",
    "冬月诗织": "portrait_shiori",
    "悠真": "portrait_yuma",
    "橘悠真": "portrait_yuma",
    "千夏": "chinatsu_talking",
    "月城千夏": "chinatsu_talking",
    "少女": "chinatsu_talking",
    "陌生少女的声音": "chinatsu_talking",
}

ENDING_STARTS = {
    980: "summer_only",
    1038: "no_fifth_person",
    1091: "other_side_of_sea",
    1196: "after_screening",
    1318: "letter_to_september",
}

FINAL_TARGETS = {
    971: 980,
    973: 1038,
    975: 1091,
    977: 1196,
}

# Full-frame CGs are deliberately attached to the single screenplay beat that
# introduces their moment.  GameScene treats `cg` as the visual layer and
# suppresses cut-out portraits, so a character depicted in the CG is never
# rendered again on top of it.
CG_BY_NODE_ID = {
    # Main story
    "p0126": "cg_main_01_first_screening", "p0137": "cg_main_02_chinatsu_in_film",
    "p0282": "cg_main_03_underground_screening", "p0552": "cg_main_04_akari_forgets",
    "p0491": "cg_main_05_rain_headphones", "p0693": "cg_main_06_seagull_cinema",
    "p0737_b2": "cg_main_07_stuck_film", "p0878": "cg_main_08_seaside_bell",
    "p0897": "cg_main_09_bell_flashback", "p0943": "cg_main_10_final_preparation",
    # Ending 1
    "p0984": "cg_end1_01_private_screening", "p0982": "cg_end1_02_childhood_summer",
    "p0989": "cg_end1_03_five_at_seaside", "p1015": "cg_end1_04_empty_seat_hairclip",
    "p1016": "cg_end1_05_peaceful_morning",
    # Ending 2
    "p1039": "cg_end2_01_last_reel", "p1054": "cg_end2_02_film_fire",
    "p1040": "cg_end2_03_girl_silhouette", "p1061": "cg_end2_04_empty_canister_photo",
    # Ending 3
    "p1096": "cg_end3_01_white_break", "p1104": "cg_end3_02_coast_in_screen",
    "p1160": "cg_end3_03_girl_farewell", "p1162": "cg_end3_04_coastal_train",
    "p1174": "cg_end3_05_white_out",
    # Ending 4
    "p0950": "cg_end4_01_restoration_table", "p1198": "cg_end4_02_private_screening",
    "p1214": "cg_end4_03_restored_seaside", "p1239": "cg_end4_04_chinatsu_appears",
    "p1262": "cg_end4_05_serious_then_light", "p1279": "cg_end4_06_gentle_goodbye",
    # Hidden ending
    "p1341": "cg_end5_01_september_rooftop", "p1343": "cg_end5_02_five_preparing",
    "p1371": "cg_end5_03_face_camera", "p1378": "cg_end5_04_last_thirty_seconds",
}

EPILOGUE_CG_BY_NODE_ID = {
    "p1021": "cg_epilogue_01_new_classroom", "p1022": "cg_epilogue_02_student_finds_script",
    "p1027": "cg_epilogue_03_close_script", "p1037": "cg_epilogue_04_fade_out",
    "p1074": "cg_epilogue_01_new_classroom", "p1075": "cg_epilogue_02_student_finds_script",
    "p1080": "cg_epilogue_03_close_script", "p1090": "cg_epilogue_04_fade_out",
    "p1179": "cg_epilogue_01_new_classroom", "p1180": "cg_epilogue_02_student_finds_script",
    "p1185": "cg_epilogue_03_close_script", "p1195": "cg_epilogue_04_fade_out",
    "p1301": "cg_epilogue_01_new_classroom", "p1302": "cg_epilogue_02_student_finds_script",
    "p1307": "cg_epilogue_03_close_script", "p1317": "cg_epilogue_04_fade_out",
}

CG_BY_NODE_ID = {**CG_BY_NODE_ID, **EPILOGUE_CG_BY_NODE_ID}

# A CG is held across the dialogue that belongs to its visual beat, rather
# than flashing for the single sentence that introduces it.  Ranges without a
# suffix only apply to the main route; the cinema-repair range is specific to
# the third investigation branch.
CG_SPANS = [
    (126, 136, "cg_main_01_first_screening", None), (137, 140, "cg_main_02_chinatsu_in_film", None),
    (281, 316, "cg_main_03_underground_screening", None), (476, 495, "cg_main_05_rain_headphones", None),
    (540, 554, "cg_main_04_akari_forgets", None), (691, 708, "cg_main_06_seagull_cinema", None),
    (733, 743, "cg_main_07_stuck_film", "_b2"), (850, 882, "cg_main_08_seaside_bell", None),
    (897, 937, "cg_main_09_bell_flashback", None), (939, 949, "cg_main_10_final_preparation", None),
    (981, 983, "cg_end1_02_childhood_summer", None), (984, 988, "cg_end1_01_private_screening", None),
    (989, 1013, "cg_end1_03_five_at_seaside", None),
    (1014, 1015, "cg_end1_04_empty_seat_hairclip", None), (1016, 1019, "cg_end1_05_peaceful_morning", None),
    (1039, 1041, "cg_end2_01_last_reel", None), (1042, 1053, "cg_end2_03_girl_silhouette", None),
    (1054, 1060, "cg_end2_02_film_fire", None), (1061, 1072, "cg_end2_04_empty_canister_photo", None),
    (1093, 1104, "cg_end3_01_white_break", None), (1105, 1118, "cg_end3_02_coast_in_screen", None),
    (1154, 1161, "cg_end3_03_girl_farewell", None), (1162, 1173, "cg_end3_04_coastal_train", None),
    (1174, 1177, "cg_end3_05_white_out", None),
    (950, 969, "cg_end4_01_restoration_table", None), (1198, 1213, "cg_end4_02_private_screening", None),
    (1214, 1238, "cg_end4_03_restored_seaside", None), (1239, 1261, "cg_end4_04_chinatsu_appears", None),
    (1262, 1278, "cg_end4_05_serious_then_light", None), (1279, 1299, "cg_end4_06_gentle_goodbye", None),
    (1341, 1359, "cg_end5_01_september_rooftop", None), (1360, 1370, "cg_end5_02_five_preparing", None),
    (1371, 1377, "cg_end5_03_face_camera", None), (1378, 1391, "cg_end5_04_last_thirty_seconds", None),
    (1021, 1022, "cg_epilogue_01_new_classroom", None), (1023, 1026, "cg_epilogue_02_student_finds_script", None),
    (1027, 1036, "cg_epilogue_03_close_script", None), (1037, 1037, "cg_epilogue_04_fade_out", None),
    (1074, 1075, "cg_epilogue_01_new_classroom", None), (1076, 1079, "cg_epilogue_02_student_finds_script", None),
    (1080, 1089, "cg_epilogue_03_close_script", None), (1090, 1090, "cg_epilogue_04_fade_out", None),
    (1179, 1180, "cg_epilogue_01_new_classroom", None), (1181, 1184, "cg_epilogue_02_student_finds_script", None),
    (1185, 1194, "cg_epilogue_03_close_script", None), (1195, 1195, "cg_epilogue_04_fade_out", None),
    (1301, 1302, "cg_epilogue_01_new_classroom", None), (1303, 1306, "cg_epilogue_02_student_finds_script", None),
    (1307, 1316, "cg_epilogue_03_close_script", None), (1317, 1317, "cg_epilogue_04_fade_out", None),
]


def cg_for_node_id(node_id: str) -> str | None:
    match = re.match(r"^p(\d{4})(.*)$", node_id)
    if not match:
        return CG_BY_NODE_ID.get(node_id)
    paragraph, suffix = int(match.group(1)), match.group(2)
    for start, end, cg, required_suffix in CG_SPANS:
        if start <= paragraph <= end and ((required_suffix is None and not suffix) or suffix == required_suffix):
            return cg
    return CG_BY_NODE_ID.get(node_id)


def clean(text: str) -> str:
    return text.strip().strip("“”")


def background_for(text: str, fallback: str | None) -> str | None:
    for needle, asset in BACKGROUND_RULES:
        if needle in text:
            return asset
    if "【背景】" in text:
        if "新学校" in text:
            return "environment_new_film_classroom_september"
        if "夕阳" in text:
            return "environment_school_rooftop_sunset"
        if "屋顶" in text or "天台" in text:
            return "environment_rooftop_september_morning"
    if "第二天清晨" in text:
        return "environment_coastal_train_summer_afternoon"
    return fallback


def bgm_for(text: str, fallback: str | None) -> str | None:
    for needle, track in BGM_BY_CHAPTER.items():
        if needle in text:
            return track
    return fallback


def parse_effect(marker: str) -> dict | None:
    m = re.search(r"【变量：([A-Za-z]+)\s*([+-]\d+)】", marker)
    if m:
        return {"kind": "variable", "key": m.group(1), "delta": int(m.group(2))}
    m = re.search(r"【线索获得\s*\d+/6：([^】]+)】", marker)
    if m:
        return {"kind": "clue", "key": m.group(1), "value": True}
    return None


class Builder:
    def __init__(self, paragraphs):
        self.paragraphs = paragraphs
        self.nodes: dict[str, dict] = {}
        self.order: list[str] = []
        self.chapter = ""
        self.scene = ""
        self.background = "environment_home_background"
        self.bgm = "M01_title_theme"
        self.speaker = ""

    @staticmethod
    def pid(index: int, suffix: str = "") -> str:
        return f"p{index:04d}{suffix}"

    def add(self, node_id: str, **payload) -> str:
        if node_id in self.nodes:
            raise ValueError(f"duplicate node ID: {node_id}")
        node = {"id": node_id, **{k: v for k, v in payload.items() if v not in (None, "", [])}}
        cg = cg_for_node_id(node_id)
        if cg:
            node["cg"] = cg
        self.nodes[node_id] = node
        self.order.append(node_id)
        return node_id

    def process_plain_range(self, start: int, end: int, branch_suffix: str = "") -> tuple[str | None, str | None]:
        first = None
        previous = None
        local_speaker = self.speaker
        for i in range(start, end):
            p = self.paragraphs[i]
            text = p.text.strip()
            style = p.style.name
            if not text:
                continue
            if style == "Choice":
                continue
            if style == "System Marker":
                # Choice effects are applied atomically when the option is picked.
                # They must not be repeated later when the branch text is read.
                continue
            if style == "Speaker":
                local_speaker = clean(text.rstrip("：:"))
                continue

            node_id = self.pid(i, branch_suffix)
            node_type = "dialogue"
            speaker = local_speaker if style == "Dialogue" else ""
            visible = clean(text) if style == "Dialogue" else text
            payload = {
                "type": node_type,
                "chapter": self.chapter,
                "scene": self.scene,
                "speaker": speaker,
                "text": visible,
                "background": self.background,
                "portrait": PORTRAITS.get(speaker),
                "bgm": self.bgm,
            }
            if style == "Stage Direction":
                payload["type"] = "direction"
                payload["effect"] = "flash" if "闪白" in text else "fade"
            current = self.add(node_id, **payload)
            if previous:
                self.nodes[previous]["next"] = current
            first = first or current
            previous = current
        return first, previous

    def choice_end(self, heading_index: int, limit: int) -> int:
        for i in range(heading_index + 1, limit):
            style = self.paragraphs[i].style.name
            if style in ("Heading 1", "Heading 2", "Heading 3"):
                return i
        return limit

    def add_choice_group(self, heading_index: int, end: int) -> tuple[str, str]:
        choice_positions = [
            i for i in range(heading_index + 1, end)
            if self.paragraphs[i].style.name == "Choice"
        ]
        join_id = self.pid(end, "_join")
        options = []
        for ordinal, pos in enumerate(choice_positions):
            branch_end = choice_positions[ordinal + 1] if ordinal + 1 < len(choice_positions) else end
            effects = []
            for j in range(pos + 1, branch_end):
                effect = parse_effect(self.paragraphs[j].text.strip())
                if effect:
                    effects.append(effect)
            if heading_index == 970:
                target = self.pid(FINAL_TARGETS[pos])
            else:
                target, last = self.process_plain_range(pos + 1, branch_end, f"_b{ordinal}")
                target = target or join_id
                if last:
                    self.nodes[last]["next"] = join_id
            options.append({
                "id": f"choice_{heading_index}_{ordinal}",
                "text": re.sub(r"^[A-ZＡ-Ｚ][：:]\s*", "", self.paragraphs[pos].text.strip()),
                "target": target,
                "effects": effects,
            })
        choice_id = self.pid(heading_index)
        self.add(
            choice_id,
            type="choice",
            chapter=self.chapter,
            scene=self.scene,
            text=self.paragraphs[heading_index].text.strip(),
            background=self.background,
            bgm=self.bgm,
            checkpoint=True,
            autosave=True,
            options=options,
        )
        self.add(join_id, type="event", chapter=self.chapter, scene=self.scene, next=None)
        return choice_id, join_id

    def build(self):
        i = START_INDEX
        previous = None
        while i < STOP_INDEX:
            p = self.paragraphs[i]
            text = p.text.strip()
            style = p.style.name
            if not text:
                i += 1
                continue

            is_choice_heading = style == "Heading 3" and (
                "选择" in text or "调查选择" in text
            )
            is_final_choice = i == 970
            if is_choice_heading or is_final_choice:
                end = self.choice_end(i, STOP_INDEX) if not is_final_choice else 980
                choice_id, join_id = self.add_choice_group(i, end)
                if previous:
                    self.nodes[previous]["next"] = choice_id
                previous = None if is_final_choice else join_id
                i = end
                continue

            if style == "Heading 1":
                self.chapter = text
                self.background = background_for(text, self.background)
                self.bgm = bgm_for(text, self.bgm)
            elif style == "Heading 2":
                self.scene = text
                self.background = background_for(text, self.background)
                self.bgm = bgm_for(text, self.bgm)
            elif "【背景】" in text:
                self.background = background_for(text, self.background)
            if style == "Speaker":
                self.speaker = clean(text.rstrip("：:"))
                i += 1
                continue
            if style == "System Marker":
                effect = parse_effect(text)
                if effect and previous:
                    self.nodes[previous].setdefault("effects", []).append(effect)
                i += 1
                continue
            if style == "Choice":
                i += 1
                continue

            node_id = self.pid(i)
            is_heading = style in ("Heading 1", "Heading 2")
            speaker = self.speaker if style == "Dialogue" else ""
            payload = {
                "type": "chapter" if is_heading else ("direction" if style == "Stage Direction" else "dialogue"),
                "chapter": self.chapter,
                "scene": self.scene,
                "speaker": speaker,
                "text": clean(text) if style == "Dialogue" else text,
                "background": self.background,
                "portrait": PORTRAITS.get(speaker),
                "bgm": self.bgm,
                "checkpoint": is_heading,
                "autosave": is_heading,
            }
            if style == "Stage Direction":
                payload["effect"] = "flash" if "闪白" in text else "fade"
            self.add(node_id, **payload)
            if previous:
                self.nodes[previous]["next"] = node_id
            previous = node_id
            i += 1

        # Final choice points directly at independent endings, never at its join.
        self.nodes.pop(self.pid(980, "_join"), None)
        if self.pid(980, "_join") in self.order:
            self.order.remove(self.pid(980, "_join"))

        # Terminate each ending before the following route.  The true route can
        # continue to the hidden route only when StoryEngine approves thresholds.
        boundaries = [(980, 1038), (1038, 1091), (1091, 1196), (1196, 1318), (1318, STOP_INDEX)]
        for start, end in boundaries:
            ids = [nid for nid in self.order if start <= int(nid[1:5]) < end and "_b" not in nid and "_join" not in nid]
            if not ids:
                continue
            last = ids[-1]
            ending_id = ENDING_STARTS[start]
            self.nodes[last].pop("next", None)
            self.nodes[last]["type"] = "ending"
            self.nodes[last]["endingID"] = ending_id
            if start == 1196:
                # The last pre-epilogue TRUE END card is paragraph 1299.
                true_card = self.pid(1299)
                if true_card in self.nodes:
                    self.nodes[true_card]["conditionalNext"] = {
                        "requiresHidden": True,
                        "target": self.pid(1318),
                        "fallback": self.nodes[true_card].get("next"),
                    }

        # Chapter index drives the gallery/chapter-select screen.
        chapters = []
        for node_id in self.order:
            node = self.nodes[node_id]
            if node.get("type") == "chapter" and node.get("checkpoint"):
                chapters.append({"id": node_id, "title": node.get("text", ""), "background": node.get("background", "")})

        data = {
            "schemaVersion": 2,
            "source": SOURCE.name,
            "sourceParagraphRange": [START_INDEX, STOP_INDEX - 1],
            "startNodeID": self.pid(56),
            "hiddenEnding": {
                "requiredClues": ["第五个杯子", "被刮去的名字", "蓝色玻璃发夹", "被海水损坏的录音", "被擦掉的编剧署名", "未冲洗的最后一张照片"],
                "thresholds": {"FaceTruth": 4, "TrustFriends": 3, "EmpathyChinatsu": 4},
            },
            "chapters": chapters,
            "nodes": self.nodes,
        }
        return data


def validate(data: dict) -> None:
    nodes = data["nodes"]
    if data["startNodeID"] not in nodes:
        raise ValueError("start node is missing")
    for node_id, node in nodes.items():
        targets = []
        if node.get("next"):
            targets.append(node["next"])
        conditional = node.get("conditionalNext") or {}
        targets.extend(x for x in (conditional.get("target"), conditional.get("fallback")) if x)
        targets.extend(option["target"] for option in node.get("options", []))
        for target in targets:
            if target not in nodes:
                raise ValueError(f"{node_id} targets missing node {target}")
    ending_ids = {node.get("endingID") for node in nodes.values() if node.get("endingID")}
    expected = set(ENDING_STARTS.values())
    if ending_ids != expected:
        raise ValueError(f"ending mismatch: expected {expected}, found {ending_ids}")


def main() -> None:
    document = Document(SOURCE)
    builder = Builder(document.paragraphs)
    data = builder.build()
    validate(data)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(f"wrote {len(data['nodes'])} nodes and {len(data['chapters'])} chapter/scene anchors to {OUTPUT}")


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"error: {exc}", file=sys.stderr)
        raise
